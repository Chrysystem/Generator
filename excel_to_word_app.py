import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import pandas as pd
from docx import Document
import os
from typing import Optional
import openpyxl
import docx
import sys
import os
import numpy
from PIL import Image, ImageTk

# === Fonctions pour les documents Word ===
def generate_convention(data_row):
    doc = Document()
    doc.add_heading('Convention de Stage', 0)
    doc.add_paragraph(f"Nom: {data_row['Nom']}")
    doc.add_paragraph(f"Formation: {data_row['Formation']}")
    doc.add_paragraph(f"Date: {data_row['Date']}")
    doc.save('Convention_Stage.docx')


def generate_emargement_from_excel(excel_file_path):
    """Génère la feuille d'émargement en utilisant le fichier Excel filtré"""
    try:
        # Charger les données du fichier Excel filtré
        df = pd.read_excel(excel_file_path)
        
        doc = Document()
        doc.add_heading('Feuille d\'Emargement', 0)
        table = doc.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Nom'
        hdr_cells[1].text = 'Formation'
        hdr_cells[2].text = 'Date'
        hdr_cells[3].text = 'Signature'

        for index, row in df.iterrows():
            row_cells = table.add_row().cells
            # Utiliser les colonnes du fichier Excel filtré
            nom = row.get('firstname', '') + ' ' + row.get('lastname', '') if 'firstname' in row and 'lastname' in row else str(row.get('Nom', 'Nom inconnu'))
            formation = row.get('course full name', 'Formation inconnue')
            date = row.get('datedebutsession', 'Date inconnue')
            
            row_cells[0].text = str(nom)
            row_cells[1].text = str(formation)
            row_cells[2].text = str(date)
            row_cells[3].text = ''

        doc.save('Feuille_Emargement.docx')
        return True
        
    except Exception as e:
        print(f"Erreur lors de la génération de la feuille d'émargement: {str(e)}")
        return False


def generate_emargement(filtered_data):
    """Ancienne fonction - maintenant dépréciée, utilise generate_emargement_from_excel"""
    # Créer un fichier Excel temporaire avec les données filtrées
    temp_excel_path = "temp_filtered_data_emargement.xlsx"
    filtered_data.to_excel(temp_excel_path, index=False)
    
    # Utiliser la nouvelle fonction
    success = generate_emargement_from_excel(temp_excel_path)
    
    # Nettoyer le fichier temporaire
    try:
        os.remove(temp_excel_path)
    except:
        pass
    
    return success


def generate_chevalets_from_excel(excel_file_path):
    """Génère les chevalets en utilisant le fichier Excel filtré pour le publipostage"""
    try:
        # Charger les données du fichier Excel filtré
        df = pd.read_excel(excel_file_path)
        
        # Chemin vers le template de chevalet
        template_path = os.path.join("Datas", "documents", "template_chevalet.docx")
        
        # Vérifier si le template existe
        if not os.path.exists(template_path):
            # Créer un template par défaut si il n'existe pas
            create_default_chevalet_template(template_path)
        
        # Créer un document Word pour le publipostage
        doc = Document(template_path)
        
        # Générer un chevalet pour chaque ligne du fichier Excel
        for index, row in df.iterrows():
            try:
                # Créer une copie du template pour chaque personne
                doc_copy = Document(template_path)
                
                # Utiliser les colonnes du fichier Excel filtré
                nom = row.get('firstname', '') + ' ' + row.get('lastname', '') if 'firstname' in row and 'lastname' in row else str(row.get('Nom', 'Nom inconnu'))
                formation = row.get('course full name', 'Formation inconnue')
                date = row.get('datedebutsession', 'Date inconnue')
                
                # Remplacer les placeholders dans le template
                replace_placeholders_in_document(doc_copy, {
                    '{{NOM}}': nom,
                    '{{FORMATION}}': formation,
                    '{{DATE}}': str(date),
                    '{{PRENOM}}': row.get('firstname', ''),
                    '{{NOM_FAMILLE}}': row.get('lastname', '')
                })
                
                # Nettoyer le nom pour le nom de fichier (enlever les caractères spéciaux)
                nom_fichier = "".join(c for c in nom if c.isalnum() or c in (' ', '-', '_')).rstrip()
                doc_copy.save(f"Chevalet_{nom_fichier}.docx")
                
            except Exception as e:
                print(f"Erreur lors de la génération du chevalet pour {nom}: {str(e)}")
                # En cas d'erreur, créer un chevalet simple
                create_simple_chevalet(nom, formation, date)
                
        return True
        
    except Exception as e:
        print(f"Erreur lors de la lecture du fichier Excel: {str(e)}")
        return False


def generate_chevalets(filtered_data):
    """Ancienne fonction - maintenant dépréciée, utilise generate_chevalets_from_excel"""
    # Créer un fichier Excel temporaire avec les données filtrées
    temp_excel_path = "temp_filtered_data.xlsx"
    filtered_data.to_excel(temp_excel_path, index=False)
    
    # Utiliser la nouvelle fonction
    success = generate_chevalets_from_excel(temp_excel_path)
    
    # Nettoyer le fichier temporaire
    try:
        os.remove(temp_excel_path)
    except:
        pass
    
    return success


def create_default_chevalet_template(template_path):
    """Crée un template de chevalet par défaut"""
    doc = Document()
    
    # Titre
    title = doc.add_heading('Chevalet de Formation', 0)
    title.alignment = 1  # Centré
    
    # Logo ou espace pour logo
    doc.add_paragraph()
    logo_para = doc.add_paragraph("LOGO TMH")
    logo_para.alignment = 1  # Centré
    
    doc.add_paragraph()
    
    # Informations du participant
    doc.add_heading('Informations du Participant', level=1)
    
    # Tableau pour les informations
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    
    # Nom
    table.cell(0, 0).text = "Nom complet:"
    table.cell(0, 1).text = "{{NOM}}"
    
    # Prénom
    table.cell(1, 0).text = "Prénom:"
    table.cell(1, 1).text = "{{PRENOM}}"
    
    # Nom de famille
    table.cell(2, 0).text = "Nom de famille:"
    table.cell(2, 1).text = "{{NOM_FAMILLE}}"
    
    doc.add_paragraph()
    
    # Informations de formation
    doc.add_heading('Informations de Formation', level=1)
    
    # Tableau pour la formation
    formation_table = doc.add_table(rows=2, cols=2)
    formation_table.style = 'Table Grid'
    
    # Formation
    formation_table.cell(0, 0).text = "Formation:"
    formation_table.cell(0, 1).text = "{{FORMATION}}"
    
    # Date
    formation_table.cell(1, 0).text = "Date:"
    formation_table.cell(1, 1).text = "{{DATE}}"
    
    doc.add_paragraph()
    
    # Espace pour signature
    doc.add_paragraph("Signature du participant:")
    doc.add_paragraph("_" * 50)
    
    # Sauvegarder le template
    os.makedirs(os.path.dirname(template_path), exist_ok=True)
    doc.save(template_path)


def replace_placeholders_in_document(doc, replacements):
    """Remplace les placeholders dans le document Word"""
    for paragraph in doc.paragraphs:
        for old_text, new_text in replacements.items():
            if old_text in paragraph.text:
                paragraph.text = paragraph.text.replace(old_text, new_text)
    
    # Remplacer aussi dans les tableaux
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for old_text, new_text in replacements.items():
                    if old_text in cell.text:
                        cell.text = cell.text.replace(old_text, new_text)


def create_simple_chevalet(nom, formation, date):
    """Crée un chevalet simple en cas d'erreur avec le template"""
    doc = Document()
    doc.add_heading(f"Chevalet: {nom}", 0)
    doc.add_paragraph(f"Formation: {formation}")
    doc.add_paragraph(f"Date: {date}")
    
    nom_fichier = "".join(c for c in nom if c.isalnum() or c in (' ', '-', '_')).rstrip()
    doc.save(f"Chevalet_{nom_fichier}.docx")


def export_filtered_excel(filtered_data, formation_name="", date_filter=""):
    """Exporte les données filtrées vers un fichier Excel pour publipostage"""
    if filtered_data.empty:
        return False
    
    # Créer un nom de fichier basé sur les filtres
    filename = f"Donnees_Filtrees"
    if formation_name:
        filename += f"_{formation_name}"
    if date_filter:
        filename += f"_{date_filter}"
    filename += ".xlsx"
    
    # Exporter vers Excel
    try:
        filtered_data.to_excel(filename, index=False)
        return filename
    except Exception as e:
        print(f"Erreur lors de l'export Excel: {e}")
        return False


# === Interface Tkinter ===
class Application(tk.Tk):
    def generate_filtered_mailmerge_without_tmhf(self):
        """Génère un Excel filtré excluant l'institution 'Toyota Material Handling France S.A.S.'
        à partir de Datas/documents/source_publipostage.xlsx
        """
        try:
            source_path = resource_path(os.path.join("Datas", "documents", "source_publipostage.xlsx"))
            if not os.path.exists(source_path):
                messagebox.showerror("Erreur", f"Fichier source introuvable:\n{source_path}")
                return

            df = pd.read_excel(source_path)

            institution_col = None
            for candidate in ["institution", "Institution", "INSTITUTION"]:
                if candidate in df.columns:
                    institution_col = candidate
                    break

            if institution_col is None:
                messagebox.showerror("Erreur", "Colonne 'institution' introuvable dans le fichier source.")
                return

            mask = ~df[institution_col].astype(str).str.contains("Toyota Material Handling France S.A.S.", na=False)
            filtered = df[mask]

            if filtered.empty:
                messagebox.showinfo("Info", "Après exclusion, aucune ligne restante.")
                return

            dest_path = resource_path(os.path.join("Datas", "documents", "source_publipostage_sans_TMHF.xlsx"))
            filtered.to_excel(dest_path, index=False)

            messagebox.showinfo(
                "Succès",
                f"Fichier généré:\n{dest_path}\n\nLignes exportées: {len(filtered)}"
            )
        except Exception as e:
            messagebox.showerror("Erreur", f"Erreur lors de la génération: {str(e)}")
    def open_convention_file(self):
        """Ouvre un fichier Word avec l'application par défaut dans le dossier Datas/documents (toujours accessible)"""
        file_path = resource_path(os.path.join("Datas", "documents", "CONVENTION-Sxx 2025-BUSSY.docx"))
        #initial_dir = resource_path(os.path.join("Datas", "documents"))
        #file_path = filedialog.askopenfilename(
        #    filetypes=[("Word Files", "*.docx")],
        #    title="Ouvrir un fichier Word",
        #    initialdir=initial_dir
        #)
        if os.path.exists(file_path):
            os.startfile(file_path)
        else:
            messagebox.showerror("Erreur", f"Fichier introuvable:\n{file_path}")

    def open_certificat_file(self):
        file_path = resource_path(os.path.join("Datas", "documents", "CERTIFICAT DE REALISATION-SXX.doc"))

        if os.path.exists(file_path):
            os.startfile(file_path)
        else:
            messagebox.showerror("Erreur", f"Fichier introuvable:\n{file_path}")




    def open_word_file(self):
        """Ouvre un fichier Word avec l'application par défaut"""
        #file_path = filedialog.askopenfilename(filetypes=[("Word Files", "*.docx")], title="Ouvrir un fichier Word")
        file_path = resource_path(os.path.join("Datas", "documents", "EMARGEMENT-SxxA-CARQUEFOU.docx"))
        if os.path.exists(file_path):
            os.startfile(file_path)
        else:
            messagebox.showerror("Erreur", f"Fichier introuvable:\n{file_path}")

    def filter_and_export_excel(self):
        """Filtre les données et exporte vers un fichier Excel pour publipostage"""
        filtered_df = self.filter_data()
        if filtered_df is None:
            return

        # Demander où sauvegarder le fichier
        save_path = os.path.join("Datas", "Log", "Log_export.xlsx")
        #save_path = filedialog.asksaveasfilename(
        #   defaultextension=".xlsx",
        #    filetypes=[("Excel Files", "*.xlsx")],
        #    title="Sauvegarder le fichier Excel filtré"
        #)
        
        #if save_path:
        try:
            filtered_df.to_excel(save_path, index=False)
            
            # Proposer d'utiliser ce fichier pour le publipostage
            messagebox.askyesno("Publipostage",f"Fichier Excel exporté avec succès!\n{len(filtered_df)} lignes exportées.\n\n")
                #f"Voulez-vous utiliser ce fichier pour le publipostage Word?\n\n"
                #f"Si oui, le fichier sera copié vers l'emplacement standard pour faciliter le publipostage.")
            
            #if response:
                # Copier le fichier vers l'emplacement standard
            excel_dest = os.path.join("Datas", "documents", "source_publipostage.xlsx")
            os.makedirs(os.path.dirname(excel_dest), exist_ok=True)
            import shutil
            shutil.copy2(save_path, excel_dest)
                
            # Créer un fichier de configuration
            config_path = os.path.join("Datas", "documents", "mailmerge_config.txt")
            colonnes = list(filtered_df.columns)
            
            with open(config_path, 'w', encoding='utf-8') as f:
                f.write(f"Fichier Excel pour publipostage: {excel_dest}\n")
                f.write(f"Nombre de lignes: {len(filtered_df)}\n")
                f.write(f"Colonnes disponibles:\n")
                for col in colonnes:
                    f.write(f"- {col}\n")
            
            messagebox.showinfo("Publipostage configuré", 
                f"Fichier configuré pour le publipostage!\n\n"
                f"Fichier copié vers: {excel_dest}\n\n"
                f"Pour utiliser dans Word:\n"
                f"1. Ouvrir Word\n"
                f"2. Publipostage > Sélectionner les destinataires > Utiliser une liste existante\n"
                f"3. Sélectionner: {excel_dest}\n\n"
                f"Colonnes disponibles: {', '.join(colonnes[:5])}{'...' if len(colonnes) > 5 else ''}")
            #else:
            messagebox.showinfo("Succès", f"Fichier Excel exporté avec succès!\n{len(filtered_df)} lignes exportées.\n\nVous pouvez maintenant utiliser ce fichier pour le publipostage dans Word.")
            
            # Ouvrir le dossier contenant le fichier
            # os.startfile(os.path.dirname(save_path))
        except Exception as e:
            messagebox.showerror("Erreur", f"Erreur lors de l'export: {str(e)}")

    def __init__(self):
        super().__init__()
        self.iconbitmap("logo-Toyota-Solo.ico")
        self.title("Rev-20250822-01")
        self.minsize(700, 400)  # Augmenté la taille minimale pour accommoder l'image

        self.file_path = None
        self.df = None

        # Thème par défaut
        self.apply_theme("Clair")

        # Créer le layout principal avec deux frames
        main_frame = ttk.Frame(self, style="TFrame")
        main_frame.pack(fill="both", expand=True, padx=10, pady=10)

        # Frame de gauche pour l'image
        left_frame = ttk.Frame(main_frame, width=200, style="TFrame")
        left_frame.pack(side="left", fill="y", padx=(0, 20))
        left_frame.pack_propagate(False)  # Empêche le frame de se redimensionner

        # Frame de droite pour les widgets
        right_frame = ttk.Frame(main_frame, style="TFrame")
        right_frame.pack(side="right", fill="both", expand=True)

        # Sélecteur de style (Clair/Sombre)
        #theme_select_frame = ttk.Frame(left_frame, style="TFrame")
        #theme_select_frame.pack(fill="x", pady=(0, 5))
        #ttk.Label(theme_select_frame, text="Style:", style="TLabel").pack(side="left")
        #self.theme_combo = ttk.Combobox(theme_select_frame, values=["Clair", "Sombre"], state="readonly", width=10)
        #self.theme_combo.current(0)
        #self.theme_combo.pack(side="left", padx=5)
        #self.theme_combo.bind("<<ComboboxSelected>>", lambda e: self.apply_theme(self.theme_combo.get()))

        # Ajouter l'image sur le côté gauche
        try:
            # Charger et redimensionner l'image
            image_path = "LogoTMH.png"
            if os.path.exists(image_path):
                # Charger l'image PNG avec PIL et la redimensionner
                pil_image = Image.open(image_path)
                # Redimensionner l'image pour qu'elle tienne dans le frame (max 150px de large)
                pil_image.thumbnail((150, 150), Image.Resampling.LANCZOS)
                self.logo_image = ImageTk.PhotoImage(pil_image)
                logo_label = ttk.Label(left_frame, image=self.logo_image, style="TLabel")
                logo_label.pack(pady=20)
            else:
                # Fallback si l'image n'existe pas
                logo_label = ttk.Label(left_frame, text="Logo TMH", style="Title.TLabel")
                logo_label.pack(pady=20)
        except Exception as e:
            # En cas d'erreur, afficher un texte à la place
            logo_label = ttk.Label(left_frame, text="Logo TMH", style="Title.TLabel")
            logo_label.pack(pady=20)

       

        # Widgets dans le frame de droite
        # Titre de l'application
        title_label = ttk.Label(right_frame, text="Générateur de Documents de Formation", style="Title.TLabel", justify="center")
        title_label.pack(pady=10)
        
        # Création d'un Notebook avec 4 onglets
        notebook = ttk.Notebook(right_frame)
        notebook.pack(fill="both", expand=True)

        tab1 = ttk.Frame(notebook, style="TFrame")
        tab2 = ttk.Frame(notebook, style="TFrame")
        tab3 = ttk.Frame(notebook, style="TFrame")
        tab4 = ttk.Frame(notebook, style="TFrame")
        tab5 = ttk.Frame(notebook, style="TFrame")



        notebook.add(tab1, text="Données & Filtre")
        notebook.add(tab2, text="Documents Word")
        notebook.add(tab3, text="À venir")
        notebook.add(tab4, text="Options")
        notebook.add(tab5, text="Settings")

        # Onglet 1: Importation, date, type, filtrer/exporter
        ttk.Button(tab1, text="Charger Fichier Excel", command=self.load_excel).pack(pady=10)
        ttk.Label(tab1, text="Date (AAAA-MM-JJ):", style="TLabel").pack()
        self.date_entry = ttk.Entry(tab1)
        self.date_entry.pack()
        ttk.Label(tab1, text="Type de Formation:").pack()
        self.formation_entry = ttk.Combobox(tab1, values=[
            "AUTOPILOT", "AUTOPILOT Niveau 2", "BASES THERMIQUE MODULES 1,2 & 3",
            "BASES TRAIGO (24V série 7 et 48R + 80V série 8)", "LEVIO STAXIO SERIE P et HC",
            "LITHIUM-ION TMHMS & TMHMI", "LSI-SSI", "OPTIO H & VECTOR R", "OSE",
            "RADIO SHUTTLE", "RRE H et RRE H2 ", "RRE H2", "TONERO 15-35 STAGE V",
            "TONERO 35-80 STAGE V", "TONERO HST STAGE V", "TRAIGO 80 Série 9 20-35",
            "TRAIGO 80 Série 9 60-80", "VECTOR A"
        ])
        self.formation_entry.pack()
        ttk.Button(tab1, width=50, text="Filtrer et Exporter Excel", command=self.filter_and_export_excel).pack(pady=10)

        # Onglet 2: Actions documents
        ttk.Button(tab2, width=50, text="Ouvrir le chevalet", command=self.open_chevalet).pack(pady=10)
        ttk.Button(tab2, width=50, text="Ouvrir feuille d'emargement", command=self.open_word_file).pack(pady=10)
        ttk.Button(tab2, width=50, text="Ouvrir convention", command=self.open_convention_file).pack(pady=10)

        # Onglet 3: Placeholder
        ttk.Label(tab3, text="Fonctionnalités à venir", style="TLabel").pack(pady=20)
        ttk.Button(tab3, width=50, text="Ouvrir Certificat", command=self.open_certificat_file).pack(pady=10)

        # Onglet 4: Publipostage
        ttk.Button(tab4, width=50, text="Sélectionner template chevalet", command=self.select_chevalet_template).pack(pady=10)
        ttk.Button(tab4, width=50, text="Sélectionner fichier Excel pour publipostage", command=self.select_excel_for_mailmerge).pack(pady=10)
        ttk.Button(tab4, width=50, text="Générer Excel sans institution TMHF", command=self.generate_filtered_mailmerge_without_tmhf).pack(pady=10)

        # Onglet 5: Settings
        ttk.Label(tab5, text="Configuration", style="TLabel").pack(pady=10)
        ttk.Button(tab5, width=50, text="Configurer fichier Excel par défaut", command=self.configure_default_excel).pack(pady=10)

    def apply_theme(self, mode: str):
        """Applique un thème clair/sombre et styles ttk."""
        try:
            style = ttk.Style()
            style.theme_use("clam")

            if str(mode).lower().startswith("sombre"):
                bg = "#1f1f1f"
                fg = "#f0f0f0"
                accent = "#d32f2f"  # Rouge TMH
                entry_bg = "#2a2a2a"
                border = "#3a3a3a"
                active = "#b71c1c"
            else:
                bg = "#f7f7f7"
                fg = "#222222"
                accent = "#d32f2f"
                entry_bg = "#ffffff"
                border = "#d0d0d0"
                active = "#b71c1c"

            self.bg_color = bg
            self.fg_color = fg
            self.accent_color = accent

            # Fond de la fenêtre principale
            self.configure(bg=bg)

            # Styles de base
            style.configure("TFrame", background=bg)
            style.configure("TLabel", background=bg, foreground=fg)
            style.configure("Title.TLabel", background=bg, foreground=fg, font=("Arial", 14, "bold"))

            style.configure(
                "TButton",
                background=border,
                foreground=fg,
                bordercolor=border,
                focusthickness=2,
                focuscolor=accent
            )
            style.map(
                "TButton",
                background=[("active", accent)],
                foreground=[("active", "white")]
            )

            style.configure(
                "Accent.TButton",
                background=accent,
                foreground="white",
                bordercolor=accent
            )
            style.map("Accent.TButton", background=[("active", active)])

            style.configure(
                "TEntry",
                fieldbackground=entry_bg,
                background=entry_bg,
                foreground=fg,
                bordercolor=border
            )
            style.configure(
                "TCombobox",
                fieldbackground=entry_bg,
                background=entry_bg,
                foreground=fg,
                bordercolor=border
            )

            self.style = style
        except Exception:
            # En cas d'erreur de style, ignorer silencieusement pour ne pas bloquer l'appli
            pass

    def open_chevalet(self):
        """Ouvre le template Word de chevalet pour le publipostage"""
        # Chemin vers le template de chevalet
        template_path = os.path.join("Datas", "documents", "template_chevalet.docx")
        
        # Vérifier si le template existe
        if not os.path.exists(template_path):
            # Créer un template par défaut si il n'existe pas
            create_default_chevalet_template(template_path)
            messagebox.showinfo("Info", "Template de chevalet par défaut créé.")
        
        try:
            # Ouvrir le template Word avec l'application par défaut
            os.startfile(template_path)
            messagebox.showinfo("Succès", 
                f"Template de chevalet ouvert!\n\n"
                f"Chemin: {template_path}\n\n"
                f"Vous pouvez maintenant utiliser ce template pour le publipostage dans Word.\n\n")
                #f"Placeholders disponibles:\n"
                #f"- {{nom}} : Nom complet\n"
                #f"- {{prenom}} : Prénom\n"
                #f"- {{nom}} : Nom de famille\n"
                #f"- {{course full name}} : Nom de la formation\n"
                #f"- {{datedebutsession}} : Date de la session")
                
        except Exception as e:
            messagebox.showerror("Erreur", f"Erreur lors de l'ouverture du template: {str(e)}")

    def select_chevalet_template(self):
        """Permet à l'utilisateur de sélectionner un template personnalisé pour les chevalets"""
        file_path = filedialog.askopenfilename(
            filetypes=[("Word Files", "*.docx")],
            title="Sélectionner un template de chevalet"
        )
        if file_path:
            # Copier le template sélectionné vers le dossier Datas/documents
            template_dest = os.path.join("Datas", "documents", "template_chevalet.docx")
            os.makedirs(os.path.dirname(template_dest), exist_ok=True)
            
            try:
                import shutil
                shutil.copy2(file_path, template_dest)
                messagebox.showinfo("Succès", f"Template de chevalet mis à jour!\n\nLe template a été copié vers: {template_dest}\n\nVous pouvez maintenant utiliser le bouton 'Ouvrir le chevalet' pour générer les chevalets avec ce nouveau template.")
            except Exception as e:
                messagebox.showerror("Erreur", f"Erreur lors de la copie du template: {str(e)}")

    def select_excel_for_mailmerge(self):
        """Permet de sélectionner un fichier Excel pour le publipostage Word"""
        file_path = filedialog.askopenfilename(
            filetypes=[("Excel Files", "*.xlsx")],
            title="Sélectionner le fichier Excel pour le publipostage"
        )
        if file_path:
            try:
                # Charger le fichier Excel sélectionné pour vérifier sa structure
                df = pd.read_excel(file_path)
                
                # Afficher les informations sur le fichier sélectionné
                colonnes = list(df.columns)
                nb_lignes = len(df)
                
                # Créer un fichier de configuration pour le publipostage
                config_path = os.path.join("Datas", "documents", "mailmerge_config.txt")
                os.makedirs(os.path.dirname(config_path), exist_ok=True)
                
                with open(config_path, 'w', encoding='utf-8') as f:
                    f.write(f"Fichier Excel sélectionné: {file_path}\n")
                    f.write(f"Nombre de lignes: {nb_lignes}\n")
                    f.write(f"Colonnes disponibles:\n")
                    for col in colonnes:
                        f.write(f"- {col}\n")
                
                # Copier le fichier Excel vers un emplacement standard
                excel_dest = os.path.join("Datas", "documents", "source_publipostage.xlsx")
                import shutil
                shutil.copy2(file_path, excel_dest)
                
                messagebox.showinfo("Succès", 
                    f"Fichier Excel sélectionné pour le publipostage!\n\n"
                    f"Fichier: {os.path.basename(file_path)}\n"
                    f"Lignes: {nb_lignes}\n"
                    f"Colonnes: {len(colonnes)}\n\n"
                    f"Le fichier a été copié vers: {excel_dest}\n\n"
                    f"Vous pouvez maintenant:\n"
                    f"1. Ouvrir Word\n"
                    f"2. Aller dans 'Publipostage' > 'Sélectionner les destinataires' > 'Utiliser une liste existante'\n"
                    f"3. Sélectionner le fichier: {excel_dest}\n\n"
                    f"Colonnes disponibles pour le publipostage:\n"
                    f"{', '.join(colonnes[:5])}{'...' if len(colonnes) > 5 else ''}")
                    
            except Exception as e:
                messagebox.showerror("Erreur", f"Erreur lors de la sélection du fichier Excel: {str(e)}")

    def configure_default_excel(self):
        """Permet de configurer le fichier Excel par défaut"""
        file_path = filedialog.askopenfilename(
            filetypes=[("Excel Files", "*.xlsx")],
            title="Sélectionner le fichier Excel par défaut"
        )
        if file_path:
            try:
                # Sauvegarder le chemin dans un fichier de configuration
                config_path = os.path.join("Datas", "config", "default_excel.txt")
                os.makedirs(os.path.dirname(config_path), exist_ok=True)
                
                with open(config_path, 'w', encoding='utf-8') as f:
                    f.write(file_path)
                
                messagebox.showinfo("Succès", 
                    f"Fichier Excel par défaut configuré!\n\n"
                    f"Chemin: {file_path}\n\n"
                    f"Ce fichier sera maintenant utilisé par le bouton 'Charger Fichier Excel'.")
                    
            except Exception as e:
                messagebox.showerror("Erreur", f"Erreur lors de la configuration: {str(e)}")

    def get_default_excel_path(self):
        """Récupère le chemin du fichier Excel par défaut depuis la configuration"""
        config_path = os.path.join("Datas", "config", "default_excel.txt")
        if os.path.exists(config_path):
            try:
                with open(config_path, 'r', encoding='utf-8') as f:
                    path = f.read().strip()
                if os.path.exists(path):
                    return path
            except:
                pass
        
        # Chemin par défaut si pas de configuration
        return resource_path(os.path.join("Datas", "Export Inscription FaceToFace pour admin Formateur175241145.xlsx"))

    def load_excel(self):
        try:
            #self.file_path = filedialog.askopenfilename(filetypes=[("Datas", "*.xlsx")])
            #ajout de fonction pour charger le fichier excel par défaut
            self.file_path = self.get_default_excel_path()
            if os.path.exists(self.file_path):
                self.df = pd.read_excel(self.file_path)
                # Afficher les informations sur les colonnes de filtrage
                info_text = "Fichier chargé avec succès!\n\n"
                
                #if 'datedebutsession' in self.df.columns:
                #    dates_uniques = self.df['datedebutsession'].dropna().unique()
                #   info_text += f"Dates disponibles: {', '.join(map(str, dates_uniques[:5]))}\n"
                #    if len(dates_uniques) > 5:
                #        info_text += f"... et {len(dates_uniques) - 5} autres dates\n"
                
                #if 'course full name' in self.df.columns:
                #    formations_uniques = self.df['course full name'].dropna().unique()
                #    info_text += f"\nFormations disponibles: {', '.join(map(str, formations_uniques[:3]))}\n"
                #    if len(formations_uniques) > 3:
                #        info_text += f"... et {len(formations_uniques) - 3} autres formations"
                
                messagebox.showinfo("Succès", info_text)
            else:
                messagebox.showerror("Erreur", "Vérifié le chemin du fichier")
        except Exception as e:
            messagebox.showerror("Erreur", f"Erreur lors du chargement du fichier Excel: {str(e)}")


    def filter_data(self) -> Optional[pd.DataFrame]:
        """Filtre les données selon les critères saisis"""
        if self.df is None:
            messagebox.showerror("Erreur", "Veuillez d'abord charger un fichier Excel.")
            return None

        date_filter = self.date_entry.get()
        formation_filter = self.formation_entry.get()

        filtered_df: pd.DataFrame = self.df.copy()
        
        # Debug: afficher le nombre de lignes initial
        initial_count = len(filtered_df)
        debug_info = f"Données initiales: {initial_count} lignes\n"

        if date_filter:
            # Filtrage plus flexible pour la date
            filtered_df = filtered_df[filtered_df['datedebutsession'].astype(str).str.contains(date_filter, na=False)]
            debug_info += f"Après filtre date '{date_filter}': {len(filtered_df)} lignes\n"
            
        if formation_filter:
            # Filtrage très flexible pour la formation (insensible à la casse et par mots-clés)
            formation_filter_lower = formation_filter.lower()
            
            # Recherche par mots-clés séparés
            mots_cles = formation_filter_lower.split()
            if len(mots_cles) > 1:
                # Si plusieurs mots-clés, chercher ceux qui contiennent TOUS les mots
                mask = pd.Series([True] * len(filtered_df), index=filtered_df.index)
                for mot in mots_cles:
                    mask = mask & filtered_df['course full name'].astype(str).str.lower().str.contains(mot, na=False)
                filtered_df = filtered_df[mask]
                debug_info += f"Après filtre formation '{formation_filter}' (recherche par mots-clés): {len(filtered_df)} lignes\n"
            else:
                # Si un seul mot, recherche simple
                filtered_df = filtered_df[filtered_df['course full name'].astype(str).str.lower().str.contains(formation_filter_lower, na=False)]
                debug_info += f"Après filtre formation '{formation_filter}' (recherche insensible à la casse): {len(filtered_df)} lignes\n"
        
        # Si aucun filtre n'est appliqué, utiliser toutes les données
        if not date_filter and not formation_filter:
            debug_info += "Aucun filtre appliqué - utilisation de toutes les données\n"

        if len(filtered_df) == 0:
            messagebox.showinfo("Aucun Résultat", f"Aucune donnée ne correspond aux critères.\n\n{debug_info}")
            return None
            
        return filtered_df



    def filter_and_generate(self):
        """Filtre les données et génère les documents Word (ancienne méthode)"""
        filtered_df = self.filter_data()
        if filtered_df is None:
            return

        generate_emargement(filtered_df)
        generate_chevalets(filtered_df)

        # Générer convention pour la première ligne filtrée
        if len(filtered_df) > 0:
            generate_convention(filtered_df.iloc[0])

        messagebox.showinfo("Succès", "Documents Word générés avec succès!")


def resource_path(relative_path):
    """Obtenir le chemin absolu vers une ressource, compatible dev et .exe PyInstaller"""
    if hasattr(sys, '_MEIPASS'):
        return os.path.join(sys._MEIPASS, relative_path)
    return os.path.join(os.path.abspath("."), relative_path)


if __name__ == "__main__":
    app = Application()
    app.mainloop()
